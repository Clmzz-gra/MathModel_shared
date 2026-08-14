# 本程序及代码是在AI工具辅助下完成的。
# AI工具名称：Trae，版本/型号：1.0（CN），开发机构/公司：字节跳动，版本发布日期：2026-08-07 到 2026-08-10。

"""
目的：
    阶段 0.1 预读 — 提取 C 题赛题 PDF 与附件1.docx 的文本，供问题理解与领域知识检索

原理：
    pypdf 提取 PDF 逐页文本；python-docx 提取 docx 段落与表格文本

输入数据：
    - C题 面向算电协同的多目标调度优化研究.pdf (原始赛题陈述)
    - 附件1.docx (原始数据说明)

输出：
    - outputs/scratch/c-problem-text.txt — 赛题全文
    - outputs/scratch/attachment1-text.txt — 附件1 文本

对应论文章节：
    阶段 0.1 启动新问题（problem-init 输入）
"""
import sys
from pathlib import Path

BASE = Path(r"e:\MathModel_pj-2026-C")
CASE_DIR = BASE / "problems" / "2026年第七届华数杯数学建模竞赛赛题" / "C题 面向算电协同的多目标调度优化研究"
OUT_DIR = BASE / "outputs" / "scratch"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# 1. PDF 文本提取
pdf_path = CASE_DIR / "C题 面向算电协同的多目标调度优化研究.pdf"
try:
    from pypdf import PdfReader
    reader = PdfReader(str(pdf_path))
    pdf_text = []
    for i, page in enumerate(reader.pages, 1):
        t = page.extract_text() or ""
        pdf_text.append(f"===== 第 {i} 页 =====\n{t}")
    (OUT_DIR / "c-problem-text.txt").write_text("\n".join(pdf_text), encoding="utf-8")
    print(f"[OK] PDF pages={len(reader.pages)} chars={sum(len(x) for x in pdf_text)}")
except Exception as e:
    print(f"[PDF FAIL] {e}", file=sys.stderr)

# 2. docx 文本提取
docx_path = CASE_DIR / "附件1.docx"
try:
    from docx import Document
    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for ti, table in enumerate(doc.tables, 1):
        parts.append(f"\n===== 表格 {ti} =====")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append(" | ".join(cells))
    (OUT_DIR / "attachment1-text.txt").write_text("\n".join(parts), encoding="utf-8")
    print(f"[OK] docx paras={len(doc.paragraphs)} tables={len(doc.tables)} chars={sum(len(x) for x in parts)}")
except Exception as e:
    print(f"[DOCX FAIL] {e}", file=sys.stderr)
